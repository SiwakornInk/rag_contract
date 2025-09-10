from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, Dict
import os
import io
from urllib.parse import quote
import oracledb
import logging

from database import OracleVectorDB
from auth import authenticate_user, create_token, get_current_user, ensure_can_upload, ensure_level, has_access, ensure_admin, hash_password, LEVEL_ORDER, ROLES, get_current_user_flexible
from embeddings import EmbeddingGenerator
from pdf_processor import PDFProcessor
from retriever import DocumentRetriever
from hybrid_retriever import HybridRetriever
from llm_handler import LLMHandler
from docx import Document as DocxDocument
import json


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Simplify for internal Docker network
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

db = OracleVectorDB()
embedder = EmbeddingGenerator()
pdf_processor = PDFProcessor()
retriever = DocumentRetriever(db, embedder)
hybrid_retriever = HybridRetriever(db, embedder)
llm_handler = LLMHandler()

logger = logging.getLogger("templates")
if not logger.handlers:
    handler = logging.StreamHandler()
    fmt = logging.Formatter('[%(asctime)s] %(levelname)s - %(name)s: %(message)s')
    handler.setFormatter(fmt)
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

class QuestionRequest(BaseModel):
    question: str
    document_filename: Optional[str] = None
    top_k: Optional[int] = 10
    
class LoginRequest(BaseModel):
    username: str
    password: str

class CreateUserRequest(BaseModel):
    username: str
    password: str
    role: str
    max_level: str

class UserOut(BaseModel):
    id: int
    username: str
    role: str
    max_level: str

class UploadResponse(BaseModel):
    status: str
    doc_id: int
    filename: str
    title: str
    total_chunks: int
    ocr_mode: str
    extraction_stats: dict
    warnings: Optional[list] = None

class PageRequest(BaseModel):
    filename: str
    page_number: int

@app.post("/auth/login")
async def login(request: LoginRequest):
    user = authenticate_user(request.username, request.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_token({
        "sub": user["username"],
        "uid": user["id"],
        "role": user["role"],
        "max_level": user["max_level"]
    })
    return {"access_token": token, "token_type": "bearer", "role": user["role"], "max_level": user["max_level"]}

@app.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return {"username": user.username, "role": user.role, "max_level": user.max_level}

@app.post("/admin/users", response_model=UserOut)
async def create_user(req: CreateUserRequest, user=Depends(get_current_user)):
    ensure_admin(user)
    if req.role not in ROLES:
        raise HTTPException(status_code=400, detail="Invalid role")
    if req.max_level not in LEVEL_ORDER:
        raise HTTPException(status_code=400, detail="Invalid level")
    # role to max_level sanity: enforce ordering (cannot assign max_level higher than SECRET anyway)
    # Insert
    try:
        with db.get_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT 1 FROM users WHERE username=:u", u=req.username)
            if cur.fetchone():
                raise HTTPException(status_code=400, detail="Username exists")
            id_var = cur.var(oracledb.NUMBER)
            cur.execute(
                "INSERT INTO users (username, password_hash, role, max_level) VALUES (:u, :p, :r, :m) RETURNING id INTO :id",
                u=req.username, p=hash_password(req.password), r=req.role, m=req.max_level, id=id_var
            )
            raw_val = id_var.getvalue()
            if isinstance(raw_val, list):
                raw_val = raw_val[0] if raw_val else None
            new_id = int(raw_val) if raw_val is not None else None
            if new_id is None:
                cur.execute("SELECT id FROM users WHERE username=:u", u=req.username)
                row = cur.fetchone()
                if not row:
                    raise HTTPException(status_code=500, detail="Failed to retrieve new user id")
                new_id = row[0]
            conn.commit()
            return {"id": new_id, "username": req.username, "role": req.role, "max_level": req.max_level}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/admin/users")
async def list_users(user=Depends(get_current_user)):
    ensure_admin(user)
    with db.get_connection() as conn:
        cur = conn.cursor()
        cur.execute("SELECT id, username, role, max_level, created_at FROM users ORDER BY id")
        rows = cur.fetchall()
        return {"users": [ {"id":r[0], "username":r[1], "role":r[2], "max_level":r[3], "created_at": r[4].isoformat() if r[4] else None } for r in rows ]}

@app.post("/upload")
async def upload_pdf(
    file: UploadFile = File(...),
    use_cloud_ocr: Optional[str] = Form("true"),
    classification: Optional[str] = Form("PUBLIC"),
    user=Depends(get_current_user)
):
    ensure_can_upload(user)
    classification = (classification or "PUBLIC").upper()
    if classification not in ["PUBLIC", "INTERNAL", "CONFIDENTIAL", "SECRET"]:
        raise HTTPException(status_code=400, detail="Invalid classification")
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    use_cloud_ocr_bool = use_cloud_ocr.lower() == "true"
    
    temp_path = f"temp_{file.filename}"
    pdf_bytes = None
    try:
        # อ่านไฟล์ทั้งหมดเพื่อเก็บใน database
        file_content = await file.read()
        pdf_bytes = file_content
        
        # เขียนไฟล์ชั่วคราวเพื่อประมวลผล
        with open(temp_path, "wb") as buffer:
            buffer.write(file_content)
        
        existing_doc = db.get_document_by_filename(file.filename)
        if existing_doc:
            raise HTTPException(status_code=400, detail="Document already exists")
        
        print(f"Processing PDF: {file.filename}")
        print(f"OCR Mode: {'Cloud (Google Vision)' if use_cloud_ocr_bool else 'Local (EasyOCR)'}")
        
        pdf_data = pdf_processor.extract_pdf_content(temp_path, use_cloud_ocr=use_cloud_ocr_bool)
        print(f"Extracted {len(pdf_data['chunks'])} chunks")
        
        extraction_stats = pdf_data.get('extraction_stats', {})
        failed_pages = extraction_stats.get('failed_pages', [])
        
        if len(pdf_data['chunks']) == 0:
            print(f"Failed to extract any content from {file.filename}")
            raise HTTPException(
                status_code=400, 
                detail=f"ไม่สามารถสกัดข้อมูลจากไฟล์ได้เลย อาจเป็นเพราะ: "
                       f"1) คุณภาพการสแกนต่ำเกินไป "
                       f"2) ไฟล์เสียหาย "
                       f"3) รูปแบบไฟล์ไม่รองรับ "
                       f"กรุณาลองใช้ Cloud OCR หรือตรวจสอบไฟล์อีกครั้ง"
            )
        
        extraction_warnings = []
        if failed_pages:
            extraction_warnings.append(f"ไม่สามารถสกัดข้อมูลจากหน้า: {', '.join(map(str, failed_pages))}")
        
        if extraction_stats.get('pages_ocr_used', 0) > 0:
            ocr_percent = (extraction_stats['pages_ocr_used'] / extraction_stats['total_pages']) * 100
            extraction_warnings.append(f"ใช้ OCR กับ {ocr_percent:.0f}% ของเอกสาร ({extraction_stats['pages_ocr_used']}/{extraction_stats['total_pages']} หน้า)")
        
        document_info = {
            'filename': file.filename,
            'original_title': pdf_data['title'],
            'abstract': pdf_data['abstract'],
            'sample_text': pdf_data.get('full_text_sample', ''),
            'document_type': pdf_data.get('document_type', 'unknown'),
            'language': pdf_data.get('metadata', {}).get('language', 'unknown'),
            'total_pages': pdf_data['total_pages']
        }
        
        print("Generating intelligent title...")
        intelligent_title = llm_handler.generate_document_title(document_info)
        print(f"Generated title: {intelligent_title}")
        
        doc_id = db.insert_document(
            filename=file.filename,
            title=intelligent_title,
            abstract=pdf_data['abstract'],
            total_pages=pdf_data['total_pages'],
            metadata={
                'original_filename': file.filename,
                'original_title': pdf_data['title'],
                'table_of_contents': pdf_data.get('table_of_contents'),
                'document_structure': pdf_data.get('document_structure'),
                'ocr_mode': 'cloud' if use_cloud_ocr_bool else 'local',
                'extraction_stats': extraction_stats,
                'classification': classification
            },
            pdf_file=pdf_bytes,
            classification=classification
        )
        print(f"Document inserted with ID: {doc_id}")
        
        for i, chunk in enumerate(pdf_data['chunks']):
            embedding = embedder.encode(chunk['text'])
            db.insert_chunk(
                doc_id=doc_id,
                chunk_text=chunk['text'],
                chunk_type=chunk['type'],
                page_number=chunk['page'],
                chunk_order=i,
                embedding=embedding,
                metadata=chunk['metadata']
            )
        
        print(f"All chunks inserted successfully")
        
        response = {
            "status": "success",
            "doc_id": doc_id,
            "filename": file.filename,
            "title": intelligent_title,
            "total_chunks": len(pdf_data['chunks']),
            "ocr_mode": 'cloud' if use_cloud_ocr_bool else 'local',
            "extraction_stats": extraction_stats
        }
        
        if extraction_warnings:
            response["warnings"] = extraction_warnings
            
        return response
    
    except Exception as e:
        print(f"Error uploading PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.post("/ask")
async def ask_question(request: QuestionRequest, user=Depends(get_current_user)):
    try:
        print(f"Question received: {request.question}")

        # If specific document specified, verify access
        if request.document_filename:
            doc = db.get_document_by_filename(request.document_filename)
            if not doc:
                raise HTTPException(status_code=404, detail="Document not found")
            if not has_access(user, doc['classification']):
                raise HTTPException(status_code=403, detail="No access to this document")

        # If no document specified and user is very low level (PUBLIC) we still allow, but all retrieval will be limited by list_documents filtering already.

        level_index = {lvl:i for i,lvl in enumerate(LEVEL_ORDER)}
        allowed_levels = [lvl for lvl in LEVEL_ORDER if level_index[lvl] <= level_index[user.max_level]]
        context_chunks = await hybrid_retriever.retrieve(
            query=request.question,
            doc_filename=request.document_filename,
            top_k=request.top_k or 15,
            allowed_levels=allowed_levels
        )
        
        print(f"Found {len(context_chunks)} relevant chunks")
        
        if not context_chunks:
            q = request.question.strip()
            list_patterns = [
                ('เอกสาร' in q or 'ไฟล์' in q or 'document' in q or 'documents' in q or 'files' in q),
                ('อะไร' in q or 'บ้าง' in q or 'ไหน' in q or 'list' in q or 'รายการ' in q)
            ]
            if all(list_patterns):
                # Fallback: list accessible documents
                docs = db.list_documents(max_level=user.max_level)
                if docs:
                    answer_lines = ["เอกสารที่คุณเข้าถึงได้ (จำกัดตามระดับสิทธิ์):"]
                    for d in docs:
                        answer_lines.append(f"- {d['title']} (ไฟล์: {d['filename']}, ระดับ: {d.get('classification','?')})")
                    return {
                        "answer": "\n".join(answer_lines),
                        "sources": [{
                            "doc_id": d['doc_id'],
                            "filename": d['filename'],
                            "title": d['title'],
                            "classification": d.get('classification')
                        } for d in docs]
                    }
                else:
                    return {"answer": "ยังไม่มีเอกสารที่คุณเข้าถึงได้", "sources": []}
            return {"answer": "ไม่พบข้อมูลที่เกี่ยวข้องกับคำถามของคุณ", "sources": []}
        
        result = llm_handler.generate_answer(request.question, context_chunks)
        
        return result
    
    except Exception as e:
        print(f"Error in ask_question: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/documents")
async def list_documents(user=Depends(get_current_user)):
    try:
        documents = db.list_documents(max_level=user.max_level)
        return {"documents": documents}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/page")
async def get_page(request: PageRequest, user=Depends(get_current_user)):
    try:
        doc = db.get_document_by_filename(request.filename)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        if not has_access(user, doc['classification']):
            raise HTTPException(status_code=403, detail="No access to this document")
        content = retriever.get_page_content(request.filename, request.page_number)
        if not content:
            raise HTTPException(status_code=404, detail="Page not found")
        
        page_text = "\n\n".join([chunk['text'] for chunk in content])
        
        return {
            "filename": request.filename,
            "page": request.page_number,
            "content": page_text,
            "chunks": content
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/document/{doc_id}/check-pdf")
async def check_pdf_exists(doc_id: int, user=Depends(get_current_user_flexible)):
    try:
        meta = db.get_document_meta(doc_id)
        if not meta or not has_access(user, meta['classification']):
            return {"exists": False, "filename": None}
        pdf_data, filename = db.get_pdf_file(doc_id)
        return {
            "exists": pdf_data is not None,
            "filename": filename if pdf_data else None
        }
    except Exception as e:
        print(f"Error checking PDF: {str(e)}")
        return {"exists": False, "filename": None}

@app.get("/document/{doc_id}/pdf")
async def get_pdf(doc_id: int, user=Depends(get_current_user_flexible)):
    try:
        print(f"Requesting PDF for doc_id: {doc_id}")
        meta = db.get_document_meta(doc_id)
        if not meta or not has_access(user, meta['classification']):
            raise HTTPException(status_code=403, detail="No access to this document")
        pdf_data, filename = db.get_pdf_file(doc_id)
        
        if not pdf_data:
            print(f"PDF not found for doc_id: {doc_id}")
            raise HTTPException(status_code=404, detail="ไม่พบไฟล์ PDF หรือเอกสารนี้อัพโหลดก่อนระบบเก็บ PDF")
        
        print(f"Returning PDF: {filename.encode('utf-8', 'ignore').decode('utf-8')} ({len(pdf_data)} bytes)")
        
        # Encode filename for Content-Disposition header
        encoded_filename = quote(filename.encode('utf-8'))
        
        return StreamingResponse(
            io.BytesIO(pdf_data),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"inline; filename*=UTF-8''{encoded_filename}",
                "Content-Type": "application/pdf",
                "Cache-Control": "no-cache",
                "Access-Control-Allow-Origin": "*",
                "X-Content-Type-Options": "nosniff"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error getting PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/document/{doc_id}/download")
async def download_pdf(doc_id: int, user=Depends(get_current_user_flexible)):
    try:
        meta = db.get_document_meta(doc_id)
        if not meta or not has_access(user, meta['classification']):
            raise HTTPException(status_code=403, detail="No access to this document")
        pdf_data, filename = db.get_pdf_file(doc_id)
        if not pdf_data:
            raise HTTPException(status_code=404, detail="ไม่พบไฟล์ PDF หรือเอกสารนี้อัพโหลดก่อนระบบเก็บ PDF")
        
        # Encode filename for Content-Disposition header
        encoded_filename = quote(filename.encode('utf-8'))
        
        return StreamingResponse(
            io.BytesIO(pdf_data),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Content-Type": "application/pdf"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error downloading PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ================= Templates Feature =================
@app.post("/templates/upload")
async def upload_template(
    file: UploadFile = File(...),
    doc_type: Optional[str] = Form(None),
    language: Optional[str] = Form(None),
    user=Depends(get_current_user)
):
    # Admin only
    ensure_admin(user)
    filename = file.filename or ''
    lower = filename.lower()
    if not (lower.endswith('.pdf') or lower.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only .pdf or .docx allowed")

    file_bytes = await file.read()
    logger.info(f"[UPLOAD] user=%s filename=%s size=%sB doc_type=%s language=%s", user.username, filename, len(file_bytes), doc_type, language)

    # Extract text
    content_text = ""
    try:
        if lower.endswith('.docx'):
            from docx import Document as _Doc
            import io as _io
            logger.info("[UPLOAD] Extracting text from DOCX")
            doc = _Doc(_io.BytesIO(file_bytes))
            lines = [p.text for p in doc.paragraphs]
            content_text = "\n".join(lines)
        elif lower.endswith('.pdf'):
            # Save temporary and use existing PDF processor
            tmp_path = f"temp_template_{os.getpid()}_{filename}"
            with open(tmp_path, "wb") as f:
                f.write(file_bytes)
            try:
                logger.info("[UPLOAD] Extracting text from PDF (Cloud OCR)")
                pdf_data = pdf_processor.extract_pdf_content(tmp_path, use_cloud_ocr=True)
                content_text = "\n".join(
                    [c.get('text', '') for c in pdf_data.get('chunks', []) if c.get('text')]
                )
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
        else:
            raise HTTPException(status_code=400, detail="Unsupported template type")
    except Exception as e:
        logger.exception("[UPLOAD] Template text extraction failed")
        raise HTTPException(status_code=500, detail=f"Template text extraction failed: {e}")

    logger.info("[UPLOAD] Extracted text length=%d chars", len(content_text))

    # Insert into DB
    template_name = os.path.splitext(os.path.basename(filename))[0]
    template_id = db.insert_template(
        name=template_name,
        original_filename=filename,
        doc_type=doc_type or 'unknown',
        language=language or 'unknown',
        file_bytes=file_bytes,
        content_text=content_text,
        created_by=user.username
    )
    logger.info("[UPLOAD] Inserted template id=%s name=%s", template_id, template_name)

    # Analyze placeholders with LLM
    fields = llm_handler.analyze_template_placeholders(content_text)
    try:
        db.update_template_fields(template_id, json.dumps(fields, ensure_ascii=False))
    except Exception as e:
        logger.exception("[UPLOAD] Failed to save fields_json")

    logger.info("[UPLOAD] Placeholder fields detected=%d names=%s", len(fields or []),
                ", ".join([f.get('placeholder_name','') for f in (fields or [])][:10]))

    return {
        "template_id": template_id,
        "name": template_name,
        "filename": filename,
        "doc_type": doc_type or 'unknown',
        "language": language or 'unknown',
        "fields": fields
    }

@app.get("/templates")
async def list_templates_api(user=Depends(get_current_user)):
    items = db.list_templates()
    return {"templates": items}

@app.get("/templates/{template_id}/fields")
async def get_template_fields(template_id: int, user=Depends(get_current_user)):
    t = db.get_template_by_id(template_id)
    if not t:
        raise HTTPException(status_code=404, detail="Template not found")
    try:
        fields = json.loads(t.get('fields_json') or '[]')
    except Exception:
        fields = []
    return {
        "template_id": t['id'],
        "name": t['name'],
        "filename": t['original_filename'],
        "fields": fields
    }

@app.post("/templates/{template_id}/generate")
async def generate_contract(template_id: int, payload: Dict, user=Depends(get_current_user)):
    t = db.get_template_by_id(template_id)
    if not t:
        raise HTTPException(status_code=404, detail="Template not found")

    values = payload or {}
    logger.info("[GENERATE] user=%s template_id=%s provided_keys=%d", user.username, template_id, len(values.keys()))
    # Validate required keys from fields_json
    try:
        expected = json.loads(t.get('fields_json') or '[]')
    except Exception:
        expected = []
    missing = []
    for f in expected:
        key = f.get('placeholder_name')
        if key and key not in values:
            missing.append(key)
    if missing:
        logger.warning("[GENERATE] Missing values for keys: %s", ", ".join(missing))
        raise HTTPException(status_code=400, detail=f"Missing values: {', '.join(missing)}")

    template_text = t.get('content_text') or ''
    logger.info("[GENERATE] Calling LLM to fill template (template_text_len=%d)", len(template_text))
    final_text = llm_handler.fill_template_with_values(template_text, values)
    logger.info("[GENERATE] Final text length=%d", len(final_text))

    # Build docx from final_text (paragraphs by newline)
    doc = DocxDocument()
    for para in final_text.split('\n'):
        doc.add_paragraph(para)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    filename_safe = f"{t['name']}_filled.docx"
    logger.info("[GENERATE] Generated DOCX bytes=%d filename=%s", out.getbuffer().nbytes, filename_safe)
    # Encode UTF-8 filename per RFC 5987 to avoid latin-1 header encoding issues
    encoded_filename = quote(filename_safe.encode('utf-8'))
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
    }
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)